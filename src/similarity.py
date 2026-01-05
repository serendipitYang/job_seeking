"""
Resume parser and BERT-based similarity matching module.
Compares job descriptions with multiple resumes and recommends the best match.
"""

import os
import re
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
import logging

logger = logging.getLogger(__name__)

# Lazy imports for heavy dependencies
_sentence_transformer = None
_docx = None


def get_sentence_transformer():
    """Lazy load sentence transformer model."""
    global _sentence_transformer
    if _sentence_transformer is None:
        try:
            from sentence_transformers import SentenceTransformer
            model_name = os.environ.get("SIMILARITY_MODEL", "all-MiniLM-L6-v2")
            logger.info(f"Loading sentence transformer model: {model_name}")
            _sentence_transformer = SentenceTransformer(model_name)
        except ImportError:
            logger.error("sentence-transformers not installed. Run: pip install sentence-transformers")
            raise
    return _sentence_transformer


def get_docx():
    """Lazy load python-docx."""
    global _docx
    if _docx is None:
        try:
            import docx
            _docx = docx
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
    return _docx


@dataclass
class ResumeInfo:
    """Parsed resume information."""
    name: str
    path: str
    description: str
    text: str
    embedding: Optional[any] = None


@dataclass
class SimilarityResult:
    """Result of similarity comparison."""
    job_title: str
    job_description: str
    scores: Dict[str, float]  # resume_name -> score
    recommended_resume: str
    recommended_score: float
    all_scores_display: str  # Formatted string for display


class ResumeParser:
    """Parses resume documents and extracts text."""

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from a .docx file."""
        docx = get_docx()
        try:
            doc = docx.Document(file_path)
            full_text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            return ""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from a PDF file."""
        try:
            import PyPDF2
            text = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            return "\n".join(text)
        except ImportError:
            logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
            return ""
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return ""

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Read text from a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading {file_path}: {e}")
            return ""

    def parse(self, file_path: str) -> str:
        """Parse resume file and return text content."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return ""

        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.docx':
            return self.parse_docx(file_path)
        elif ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext == '.txt':
            return self.parse_txt(file_path)
        else:
            logger.error(f"Unsupported file format: {ext}")
            return ""


class SimilarityMatcher:
    """BERT-based similarity matcher for comparing resumes with job descriptions."""

    def __init__(self, config: Dict):
        self.config = config
        self.parser = ResumeParser()
        self.resumes: Dict[str, ResumeInfo] = {}
        self._model = None
        self._embeddings_computed = False

    def _get_model(self):
        """Lazy load the model."""
        if self._model is None:
            self._model = get_sentence_transformer()
        return self._model

    def load_resumes(self, base_path: str = ".") -> None:
        """Load and parse all configured resumes."""
        resumes_config = self.config.get("resumes", {})

        for name, info in resumes_config.items():
            path = info.get("path", "")
            full_path = os.path.join(base_path, path) if not os.path.isabs(path) else path

            if os.path.exists(full_path):
                text = self.parser.parse(full_path)
                if text:
                    self.resumes[name] = ResumeInfo(
                        name=name,
                        path=full_path,
                        description=info.get("description", name),
                        text=text,
                    )
                    logger.info(f"Loaded resume: {name} ({len(text)} chars)")
                else:
                    logger.warning(f"Empty resume content: {full_path}")
            else:
                logger.warning(f"Resume file not found: {full_path}")

    def compute_embeddings(self) -> None:
        """Pre-compute embeddings for all resumes."""
        if self._embeddings_computed:
            return

        model = self._get_model()
        for name, resume in self.resumes.items():
            # Use the first 512 words for embedding (model limitation)
            text = " ".join(resume.text.split()[:512])
            resume.embedding = model.encode(text, convert_to_tensor=True)
            logger.info(f"Computed embedding for: {name}")

        self._embeddings_computed = True

    def compute_similarity(self, job_title: str, job_description: str) -> SimilarityResult:
        """
        Compute similarity between job and all resumes.
        Returns the best matching resume recommendation.
        """
        if not self.resumes:
            logger.warning("No resumes loaded")
            return SimilarityResult(
                job_title=job_title,
                job_description=job_description,
                scores={},
                recommended_resume="N/A",
                recommended_score=0.0,
                all_scores_display="No resumes loaded",
            )

        # Ensure embeddings are computed
        self.compute_embeddings()

        model = self._get_model()

        # Combine title and description for better matching
        job_text = f"{job_title}. {job_description}"
        # Limit to 512 words
        job_text = " ".join(job_text.split()[:512])

        job_embedding = model.encode(job_text, convert_to_tensor=True)

        # Calculate similarity with each resume
        from sentence_transformers import util
        scores = {}
        for name, resume in self.resumes.items():
            if resume.embedding is not None:
                similarity = util.cos_sim(job_embedding, resume.embedding).item()
                scores[name] = round(similarity, 4)

        # Find best match
        if scores:
            best_resume = max(scores, key=scores.get)
            best_score = scores[best_resume]
        else:
            best_resume = "N/A"
            best_score = 0.0

        # Format scores for display
        scores_display = " | ".join([
            f"{self.resumes[name].description}: {score:.2%}"
            for name, score in sorted(scores.items(), key=lambda x: -x[1])
        ])

        return SimilarityResult(
            job_title=job_title,
            job_description=job_description,
            scores=scores,
            recommended_resume=best_resume,
            recommended_score=best_score,
            all_scores_display=scores_display,
        )

    def batch_compute_similarity(self, jobs: List[Dict]) -> List[SimilarityResult]:
        """
        Compute similarity for a batch of jobs.
        More efficient than computing one at a time.
        """
        results = []
        for job in jobs:
            result = self.compute_similarity(
                job.get("title", ""),
                job.get("description", "")
            )
            results.append(result)
        return results


def create_matcher(config: Dict, base_path: str = ".") -> SimilarityMatcher:
    """Factory function to create and initialize a similarity matcher."""
    matcher = SimilarityMatcher(config)
    matcher.load_resumes(base_path)
    return matcher
